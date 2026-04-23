#!/usr/bin/env python3
"""
Generate a comprehensive Word document for Meeting Transcript Storage System
Medium Scale Approach (100-1000 meetings)
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_bullet_point(doc, text, level=0):
    """Add a bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Light gray background effect
    return p

def create_document():
    """Create the comprehensive meeting transcript storage guide"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Meeting Transcript Storage System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Medium Scale Implementation Guide (100-1000 Meetings)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph()
    
    # Executive Summary
    add_heading(doc, 'Executive Summary', 1)
    add_paragraph(doc, 
        'This document outlines a comprehensive system for storing, organizing, and analyzing '
        'meeting transcripts using a Git repository with structured Markdown files and a JSON '
        'metadata index. This approach is optimized for 100-1000 meetings and provides efficient '
        'AI-powered analysis capabilities.')
    
    doc.add_paragraph()
    
    # System Overview
    add_heading(doc, '1. System Overview', 1)
    
    add_heading(doc, '1.1 Architecture', 2)
    add_paragraph(doc, 'The system uses a hybrid approach combining:')
    add_bullet_point(doc, 'Git repository for version control and backup')
    add_bullet_point(doc, 'Markdown files with YAML frontmatter for structured data')
    add_bullet_point(doc, 'JSON index for fast metadata queries')
    add_bullet_point(doc, 'Chronological folder organization (year/month)')
    
    doc.add_paragraph()
    
    add_heading(doc, '1.2 Key Benefits', 2)
    add_bullet_point(doc, 'Version Control: Track all changes and edits over time')
    add_bullet_point(doc, 'AI-Friendly: Easy for AI assistants to read and analyze')
    add_bullet_point(doc, 'Human-Readable: Browse and edit files directly')
    add_bullet_point(doc, 'Scalable: Handles 100-1000 meetings efficiently')
    add_bullet_point(doc, 'Searchable: Full-text search and metadata filtering')
    add_bullet_point(doc, 'Portable: No database dependencies, works offline')
    
    doc.add_page_break()
    
    # Repository Structure
    add_heading(doc, '2. Repository Structure', 1)
    
    add_paragraph(doc, 'Recommended folder organization:')
    add_code_block(doc, '''meeting-transcripts/
├── .git/
├── meetings/
│   ├── 2026/
│   │   ├── 04/
│   │   │   ├── 2026-04-22_planning_team-alpha.md
│   │   │   ├── 2026-04-23_standup_team-alpha.md
│   │   │   └── 2026-04-25_review_product.md
│   │   └── 05/
│   │       └── ...
│   └── 2025/
│       └── ...
├── index/
│   ├── meetings.json
│   └── tags.json
├── templates/
│   └── meeting-template.md
├── scripts/
│   ├── add_meeting.py
│   └── update_index.py
└── README.md''')
    
    doc.add_paragraph()
    
    # File Format
    add_heading(doc, '3. File Format Specification', 1)
    
    add_heading(doc, '3.1 Markdown with YAML Frontmatter', 2)
    add_paragraph(doc, 'Each meeting transcript is stored as a Markdown file with YAML frontmatter for metadata:')
    
    add_code_block(doc, '''---
date: 2026-04-22
time: "14:00-15:00"
duration_minutes: 60
type: planning
attendees:
  - alice@company.com
  - bob@company.com
  - charlie@company.com
facilitator: alice@company.com
project: product-roadmap
tags:
  - q2-planning
  - api-redesign
  - mobile-app
status: final
---

# Meeting: Q2 Product Roadmap Planning

## Quick Summary
Discussed Q2 priorities. Decided to focus on API 
redesign first. Mobile updates postponed to Q3.

## Transcription

**[00:00] Alice:** Let's start with the Q2 priorities...

**[00:02] Bob:** I think we should focus on the API 
redesign first...

[Transcription continues...]

## Extracted Insights

### Decisions
- [DECISION] Focus on API redesign for Q2
  Owner: @bob - Date: 2026-04-22

### Action Items
- [ ] Create API redesign spec
  Owner: @bob - Due: 2026-04-29

### Blockers
- Need design approval before starting API work''')
    
    doc.add_page_break()
    
    add_heading(doc, '3.2 File Naming Convention', 2)
    add_paragraph(doc, 'Format: YYYY-MM-DD_meeting-type_team-or-project.md', bold=True)
    add_paragraph(doc, 'Examples:')
    add_bullet_point(doc, '2026-04-22_planning_team-alpha.md')
    add_bullet_point(doc, '2026-04-23_standup_engineering.md')
    add_bullet_point(doc, '2026-04-25_review_product-roadmap.md')
    add_bullet_point(doc, '2026-04-30_retrospective_q1.md')
    
    add_paragraph(doc, 'Benefits of this naming convention:')
    add_bullet_point(doc, 'Chronological sorting by default', level=1)
    add_bullet_point(doc, 'Easy to find meetings by date', level=1)
    add_bullet_point(doc, 'Descriptive at a glance', level=1)
    add_bullet_point(doc, 'No spaces (git-friendly)', level=1)
    add_bullet_point(doc, 'Consistent and predictable', level=1)
    
    doc.add_paragraph()
    
    # Metadata Index
    add_heading(doc, '4. Metadata Index System', 1)
    
    add_heading(doc, '4.1 meetings.json Structure', 2)
    add_paragraph(doc, 'The JSON index enables fast queries without reading all files:')
    
    add_code_block(doc, '''{
  "meetings": [
    {
      "id": "2026-04-22_planning_team-alpha",
      "file": "meetings/2026/04/2026-04-22_planning_team-alpha.md",
      "date": "2026-04-22",
      "time": "14:00-15:00",
      "duration_minutes": 60,
      "type": "planning",
      "title": "Q2 Product Roadmap Planning",
      "attendees": [
        "alice@company.com",
        "bob@company.com",
        "charlie@company.com"
      ],
      "facilitator": "alice@company.com",
      "project": "product-roadmap",
      "tags": ["q2-planning", "api-redesign", "mobile-app"],
      "has_decisions": true,
      "has_action_items": true,
      "has_blockers": true,
      "action_item_count": 2,
      "decision_count": 2,
      "word_count": 1250,
      "status": "final"
    }
  ],
  "last_updated": "2026-04-22T20:00:00Z"
}''')
    
    doc.add_paragraph()
    
    add_heading(doc, '4.2 Index Benefits', 2)
    add_bullet_point(doc, 'Fast filtering by date, type, attendees, or tags')
    add_bullet_point(doc, 'Quick statistics without reading full files')
    add_bullet_point(doc, 'Enables dashboard and reporting tools')
    add_bullet_point(doc, 'Supports complex queries across meetings')
    
    doc.add_page_break()
    
    # Workflow
    add_heading(doc, '5. Implementation Workflow', 1)
    
    add_heading(doc, '5.1 Adding a New Meeting', 2)
    add_paragraph(doc, 'Step-by-step process:')
    
    p = doc.add_paragraph('', style='List Number')
    p.add_run('Meeting occurs and is auto-transcribed').bold = True
    
    p = doc.add_paragraph('Export transcription as TXT file', style='List Number')
    
    p = doc.add_paragraph('', style='List Number')
    p.add_run('Add YAML frontmatter with metadata:').bold = True
    add_bullet_point(doc, 'Date, time, duration', level=1)
    add_bullet_point(doc, 'Meeting type (planning, standup, review, etc.)', level=1)
    add_bullet_point(doc, 'Attendees and facilitator', level=1)
    add_bullet_point(doc, 'Project/team identifier', level=1)
    add_bullet_point(doc, 'Relevant tags', level=1)
    
    p = doc.add_paragraph('', style='List Number')
    p.add_run('Save as Markdown file').bold = True
    p.add_run(' in appropriate year/month folder')
    
    p = doc.add_paragraph('', style='List Number')
    p.add_run('Update meetings.json index').bold = True
    p.add_run(' (manually or via script)')
    
    p = doc.add_paragraph('', style='List Number')
    p.add_run('Commit to Git repository').bold = True
    
    doc.add_paragraph()
    
    add_heading(doc, '5.2 Automation Script Example', 2)
    add_paragraph(doc, 'Python script to automate the process:')
    
    add_code_block(doc, '''import os
import json
from datetime import datetime

def add_meeting(txt_file, metadata):
    """Convert TXT transcript to MD with frontmatter"""
    
    # Read transcript
    with open(txt_file, 'r') as f:
        transcript = f.read()
    
    # Generate filename
    date = metadata['date']
    meeting_type = metadata['type']
    project = metadata['project']
    filename = f"{date}_{meeting_type}_{project}.md"
    
    # Create markdown with frontmatter
    md_content = f"""---
date: {metadata['date']}
time: {metadata['time']}
type: {metadata['type']}
attendees: {metadata['attendees']}
project: {metadata['project']}
tags: {metadata['tags']}
---

# Meeting: {metadata['title']}

## Transcription

{transcript}
"""
    
    # Save to appropriate folder
    year, month = date.split('-')[:2]
    folder = f"meetings/{year}/{month}"
    os.makedirs(folder, exist_ok=True)
    
    filepath = f"{folder}/{filename}"
    with open(filepath, 'w') as f:
        f.write(md_content)
    
    # Update index
    update_index(filepath, metadata)
    
    return filepath''')
    
    doc.add_page_break()
    
    # AI Analysis
    add_heading(doc, '6. AI-Powered Analysis', 1)
    
    add_heading(doc, '6.1 Retrieval Strategies', 2)
    add_paragraph(doc, 'The system supports multiple query patterns:')
    
    add_paragraph(doc, '1. Single Meeting Analysis', bold=True)
    add_bullet_point(doc, 'Direct file read for detailed analysis', level=1)
    add_bullet_point(doc, 'Extract decisions, action items, blockers', level=1)
    
    add_paragraph(doc, '2. Metadata Queries', bold=True)
    add_bullet_point(doc, 'Search index by date, type, attendees, tags', level=1)
    add_bullet_point(doc, 'Filter meetings by criteria', level=1)
    add_bullet_point(doc, 'Generate statistics and reports', level=1)
    
    add_paragraph(doc, '3. Full-Text Search', bold=True)
    add_bullet_point(doc, 'Search across all transcripts', level=1)
    add_bullet_point(doc, 'Find specific topics or keywords', level=1)
    add_bullet_point(doc, 'Identify patterns and trends', level=1)
    
    add_paragraph(doc, '4. Time-Range Analysis', bold=True)
    add_bullet_point(doc, 'Analyze meetings within date ranges', level=1)
    add_bullet_point(doc, 'Track progress over time', level=1)
    add_bullet_point(doc, 'Identify seasonal patterns', level=1)
    
    add_paragraph(doc, '5. Participant Analysis', bold=True)
    add_bullet_point(doc, 'Track individual participation', level=1)
    add_bullet_point(doc, 'Analyze contribution patterns', level=1)
    add_bullet_point(doc, 'Identify engagement levels', level=1)
    
    doc.add_paragraph()
    
    add_heading(doc, '6.2 Insights Generated', 2)
    add_paragraph(doc, 'AI can automatically extract:')
    add_bullet_point(doc, 'Action Items: Tasks, owners, deadlines, priorities')
    add_bullet_point(doc, 'Decisions: What was decided, by whom, rationale')
    add_bullet_point(doc, 'Blockers/Risks: Issues, impacts, mitigation plans')
    add_bullet_point(doc, 'Key Topics: Main discussion themes and time allocation')
    add_bullet_point(doc, 'Sentiment: Tone, enthusiasm, concerns expressed')
    add_bullet_point(doc, 'Follow-ups: Open questions and next steps')
    add_bullet_point(doc, 'Patterns: Recurring themes across meetings')
    add_bullet_point(doc, 'Metrics: Meeting effectiveness, completion rates')
    
    doc.add_page_break()
    
    # Best Practices
    add_heading(doc, '7. Best Practices', 1)
    
    add_heading(doc, '7.1 Metadata Quality', 2)
    add_bullet_point(doc, 'Always include date, time, and attendees')
    add_bullet_point(doc, 'Use consistent meeting type classifications')
    add_bullet_point(doc, 'Tag meetings with relevant project/topic tags')
    add_bullet_point(doc, 'Identify facilitator and note-taker')
    add_bullet_point(doc, 'Mark status (draft, final, reviewed)')
    
    doc.add_paragraph()
    
    add_heading(doc, '7.2 Transcription Quality', 2)
    add_bullet_point(doc, 'Include speaker identification')
    add_bullet_point(doc, 'Add timestamps for key moments')
    add_bullet_point(doc, 'Correct obvious transcription errors')
    add_bullet_point(doc, 'Note when audio was unclear')
    add_bullet_point(doc, 'Include context for references')
    
    doc.add_paragraph()
    
    add_heading(doc, '7.3 Git Workflow', 2)
    add_bullet_point(doc, 'Commit new meetings promptly')
    add_bullet_point(doc, 'Use descriptive commit messages')
    add_bullet_point(doc, 'Update index with each new meeting')
    add_bullet_point(doc, 'Regular backups to remote repository')
    add_bullet_point(doc, 'Branch for major reorganizations')
    
    doc.add_paragraph()
    
    add_heading(doc, '7.4 Privacy and Security', 2)
    add_bullet_point(doc, 'Use private repository for sensitive meetings')
    add_bullet_point(doc, 'Redact confidential information if needed')
    add_bullet_point(doc, 'Control access via Git permissions')
    add_bullet_point(doc, 'Consider encryption for highly sensitive data')
    add_bullet_point(doc, 'Follow company data retention policies')
    
    doc.add_page_break()
    
    # Maintenance
    add_heading(doc, '8. System Maintenance', 1)
    
    add_heading(doc, '8.1 Regular Tasks', 2)
    add_paragraph(doc, 'Weekly:', bold=True)
    add_bullet_point(doc, 'Add new meeting transcripts', level=1)
    add_bullet_point(doc, 'Update index file', level=1)
    add_bullet_point(doc, 'Commit changes to Git', level=1)
    
    add_paragraph(doc, 'Monthly:', bold=True)
    add_bullet_point(doc, 'Review and clean up tags', level=1)
    add_bullet_point(doc, 'Verify index accuracy', level=1)
    add_bullet_point(doc, 'Archive old meetings if needed', level=1)
    
    add_paragraph(doc, 'Quarterly:', bold=True)
    add_bullet_point(doc, 'Analyze meeting patterns and trends', level=1)
    add_bullet_point(doc, 'Review and update templates', level=1)
    add_bullet_point(doc, 'Optimize folder structure if needed', level=1)
    
    doc.add_paragraph()
    
    add_heading(doc, '8.2 Scaling Considerations', 2)
    add_paragraph(doc, 'As the system grows beyond 1000 meetings:')
    add_bullet_point(doc, 'Consider adding SQLite database for complex queries')
    add_bullet_point(doc, 'Implement full-text search indexing (e.g., Elasticsearch)')
    add_bullet_point(doc, 'Archive older meetings to separate repository')
    add_bullet_point(doc, 'Add automated index validation')
    add_bullet_point(doc, 'Consider dedicated search/analysis tools')
    
    doc.add_page_break()
    
    # Tools and Resources
    add_heading(doc, '9. Recommended Tools', 1)
    
    add_heading(doc, '9.1 Essential Tools', 2)
    add_bullet_point(doc, 'Git Client: GitHub Desktop, GitKraken, or command line')
    add_bullet_point(doc, 'Text Editor: VS Code (excellent Markdown support)')
    add_bullet_point(doc, 'Transcription: Otter.ai, Microsoft Teams, Zoom, Google Meet')
    add_bullet_point(doc, 'Python: For automation scripts')
    
    doc.add_paragraph()
    
    add_heading(doc, '9.2 Optional Enhancements', 2)
    add_bullet_point(doc, 'ripgrep: Fast full-text search')
    add_bullet_point(doc, 'jq: JSON query tool for index manipulation')
    add_bullet_point(doc, 'Obsidian: Advanced Markdown editor with linking')
    add_bullet_point(doc, 'GitHub Actions: Automated index updates')
    
    doc.add_page_break()
    
    # Getting Started
    add_heading(doc, '10. Getting Started Checklist', 1)
    
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('☐ Create Git repository').bold = True
    
    p = doc.add_paragraph('☐ Set up folder structure (meetings/, index/, templates/)', style='List Bullet')
    
    p = doc.add_paragraph('☐ Create meeting template file', style='List Bullet')
    
    p = doc.add_paragraph('☐ Initialize meetings.json index', style='List Bullet')
    
    p = doc.add_paragraph('☐ Document naming conventions in README', style='List Bullet')
    
    p = doc.add_paragraph('☐ Set up transcription tool (Otter.ai, Teams, etc.)', style='List Bullet')
    
    p = doc.add_paragraph('☐ Create automation script for adding meetings', style='List Bullet')
    
    p = doc.add_paragraph('☐ Test workflow with first meeting', style='List Bullet')
    
    p = doc.add_paragraph('☐ Configure Git remote (GitHub, GitLab, etc.)', style='List Bullet')
    
    p = doc.add_paragraph('☐ Train team on workflow', style='List Bullet')
    
    doc.add_paragraph()
    
    # Conclusion
    add_heading(doc, '11. Conclusion', 1)
    add_paragraph(doc, 
        'This medium-scale approach provides a robust, scalable system for managing 100-1000 '
        'meeting transcripts. The combination of Git version control, structured Markdown files, '
        'and JSON indexing creates an efficient foundation for AI-powered analysis while remaining '
        'simple enough for human use.')
    
    doc.add_paragraph()
    add_paragraph(doc, 
        'The system balances automation with flexibility, allowing teams to capture comprehensive '
        'meeting intelligence without complex infrastructure. As your needs grow, the architecture '
        'supports seamless scaling to larger datasets and more sophisticated analysis tools.')
    
    doc.add_paragraph()
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Document Version: 1.0')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].italic = True
    
    footer = doc.add_paragraph('Created: April 22, 2026')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].italic = True
    
    return doc

if __name__ == '__main__':
    print("Creating Meeting Transcript Storage System Guide...")
    doc = create_document()
    
    filename = 'Meeting_Transcript_Storage_System_Guide.docx'
    doc.save(filename)
    
    print(f"[SUCCESS] Document created successfully: {filename}")
    print(f"[SUCCESS] Location: {os.path.abspath(filename)}")

# Made with Bob
