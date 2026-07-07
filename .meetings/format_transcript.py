import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the JSON file
with open('5-11-2026 meeting with Bunge.txt', 'r', encoding='utf-8') as f:
    data = json.load(f)

# Create a new Word document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Add title
title = doc.add_heading('Meeting Transcript - Bunge', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add metadata section
doc.add_heading('Meeting Information', 1)
metadata_table = doc.add_table(rows=3, cols=2)
metadata_table.style = 'Light Grid Accent 1'

metadata_table.rows[0].cells[0].text = 'Date:'
metadata_table.rows[0].cells[1].text = 'May 11, 2026'
metadata_table.rows[1].cells[0].text = 'Last Updated:'
metadata_table.rows[1].cells[1].text = data['last_updated']
metadata_table.rows[2].cells[0].text = 'Total Segments:'
metadata_table.rows[2].cells[1].text = str(data['total_segments'])

# Add page break
doc.add_page_break()

# Add transcript section
doc.add_heading('Transcript', 1)

# Process each segment
for segment in data['segments']:
    timestamp = segment['display_time']
    text = segment['text']
    
    # Add timestamp as bold
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    
    run_time = p.add_run(f'[{timestamp}] ')
    run_time.bold = True
    run_time.font.size = Pt(10)
    
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)

# Save the document
doc.save('5-11-2026_Bunge_Meeting_Summary.docx')
print('Word document created successfully: 5-11-2026_Bunge_Meeting_Summary.docx')

# Made with Bob
