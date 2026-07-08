"""Text extraction service.

Accepts a filename and raw bytes, returns a clean UTF-8 string.

Supported formats
-----------------
.txt / .md          — UTF-8 plain text
.docx               — Word Open XML (python-docx), includes paragraphs + tables
.doc                — Legacy binary Word (docx2txt fallback)
.pdf                — PDF (pdfminer.six for layout-aware extraction,
                      falls back to pypdf if pdfminer is unavailable)
.eml                — RFC-2822 email files (stdlib email module)
.msg                — Outlook binary message files (extract-msg)
.pst                — Outlook Personal Storage Table (libratom / libpff);
                      extracts all message bodies from all folders
Any other type      — Attempted as UTF-8; raises ValueError on failure.
"""

from __future__ import annotations

import email
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def extract_text(filename: str, data: bytes) -> str:
    """Return plain text extracted from *data* based on *filename* extension.

    Args:
        filename: Original filename (used to detect format via extension).
        data:     Raw file bytes.

    Returns:
        A clean UTF-8 string with all extracted text content.

    Raises:
        ValueError: If the file type is unsupported and content is not UTF-8.
    """
    suffix = Path(filename).suffix.lower()

    if suffix in (".txt", ".md"):
        return _extract_plain(data)

    if suffix == ".docx":
        return _extract_docx(data)

    if suffix == ".doc":
        return _extract_doc(data)

    if suffix == ".pdf":
        return _extract_pdf(data)

    if suffix == ".eml":
        return _extract_eml(data)

    if suffix == ".msg":
        return _extract_msg(data)

    if suffix == ".pst":
        return _extract_pst(data)

    # Fallback: attempt UTF-8 decode for unknown types
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError(
            f"Unsupported file type '{suffix}' and content is not valid UTF-8."
        ) from exc


# ---------------------------------------------------------------------------
# Format handlers
# ---------------------------------------------------------------------------


def _extract_plain(data: bytes) -> str:
    """Decode plain text; try UTF-8 then fall back to latin-1."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


def _extract_docx(data: bytes) -> str:
    """Extract text from a .docx file using python-docx.

    Extracts:
    - All paragraph text (in document order)
    - All table cell text
    - Core document properties (title, author, subject) as a header block
    """
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []

    # Document properties (metadata)
    props = doc.core_properties
    meta: list[str] = []
    if props.title:
        meta.append(f"Title: {props.title}")
    if props.author:
        meta.append(f"Author: {props.author}")
    if props.subject:
        meta.append(f"Subject: {props.subject}")
    if props.created:
        meta.append(f"Created: {props.created}")
    if meta:
        parts.append("\n".join(meta))
        parts.append("")  # blank separator

    # Body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _extract_doc(data: bytes) -> str:
    """Extract text from a legacy .doc binary file using docx2txt.

    docx2txt can handle some binary .doc files, but for heavily formatted
    or very old .doc files the output may be partial.
    """
    try:
        import docx2txt

        # docx2txt.process() accepts a file-like object
        text = docx2txt.process(io.BytesIO(data))
        if text and text.strip():
            return text
    except Exception as exc:
        logger.warning("docx2txt failed for .doc file: %s — trying UTF-8 fallback", exc)

    # Last resort: raw UTF-8/latin-1 decode (will include binary noise but
    # recovers readable text from simple .doc files)
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return data.decode("latin-1", errors="replace")


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF.

    Tries pdfminer.six first (better layout handling for multi-column, tables).
    Falls back to pypdf if pdfminer is not installed or fails.
    """
    # --- pdfminer.six (preferred) ---
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract

        text = pdfminer_extract(io.BytesIO(data))
        if text and text.strip():
            return text
    except ImportError:
        logger.debug("pdfminer.six not installed; falling back to pypdf")
    except Exception as exc:
        logger.warning("pdfminer extraction failed: %s — falling back to pypdf", exc)

    # --- pypdf fallback ---
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception as exc:
        raise ValueError(f"PDF text extraction failed: {exc}") from exc


def _extract_eml(data: bytes) -> str:
    """Extract text from an RFC-2822 .eml email file using Python's stdlib.

    Extracts: From, To, CC, Date, Subject headers + all text/* body parts.
    Skips attachments (non-text parts).
    """
    msg = email.message_from_bytes(data)
    parts: list[str] = []

    # Headers
    for header in ("From", "To", "CC", "Date", "Subject"):
        value = msg.get(header, "")
        if value:
            parts.append(f"{header}: {value}")
    if parts:
        parts.append("")  # blank line after headers

    # Body parts
    if msg.is_multipart():
        for part in msg.walk():
            ct = part.get_content_type()
            if ct == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    charset = part.get_content_charset() or "utf-8"
                    parts.append(payload.decode(charset, errors="replace"))
            elif ct == "text/html":
                # Strip HTML tags for a plain-text approximation
                payload = part.get_payload(decode=True)
                if payload:
                    charset = part.get_content_charset() or "utf-8"
                    html = payload.decode(charset, errors="replace")
                    parts.append(_strip_html(html))
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            charset = msg.get_content_charset() or "utf-8"
            parts.append(payload.decode(charset, errors="replace"))

    return "\n".join(parts)


def _extract_msg(data: bytes) -> str:
    """Extract text from an Outlook .msg binary file using extract-msg.

    Extracts: sender, recipients, date, subject, and plain-text body.
    """
    try:
        import extract_msg  # extract-msg package

        with extract_msg.openMsg(io.BytesIO(data)) as msg:
            parts: list[str] = []

            if msg.sender:
                parts.append(f"From: {msg.sender}")
            if msg.to:
                parts.append(f"To: {msg.to}")
            if msg.cc:
                parts.append(f"CC: {msg.cc}")
            if msg.date:
                parts.append(f"Date: {msg.date}")
            if msg.subject:
                parts.append(f"Subject: {msg.subject}")
            if parts:
                parts.append("")  # blank separator

            body = msg.body or ""
            if body.strip():
                parts.append(body)
            elif msg.htmlBody:
                # Fall back to HTML body stripped of tags
                parts.append(_strip_html(msg.htmlBody))

            return "\n".join(parts)
    except Exception as exc:
        raise ValueError(f".msg extraction failed: {exc}") from exc


def _extract_pst(data: bytes) -> str:
    """Extract all email text from an Outlook .pst file using pypff.

    pypff is a standalone Python binding for libpff that builds cleanly
    on Python 3.12.  Install with:
        sudo apt install -y libpff-dev
        pip install pypff

    Walks every folder recursively and extracts the subject, sender, date,
    and plain-text (or HTML-stripped) body of every message.
    """
    try:
        import pypff  # pip install pypff (requires: sudo apt install libpff-dev)
    except ImportError as exc:
        raise ValueError(
            ".pst extraction requires pypff.\n"
            "Install it with:\n"
            "  sudo apt install -y libpff-dev\n"
            "  pip install pypff"
        ) from exc

    import os
    import tempfile

    def _walk_folder(folder: "pypff.folder") -> list[str]:
        """Recursively walk a pypff folder and return extracted message texts."""
        results: list[str] = []
        folder_name = folder.name or "Unknown Folder"

        for i in range(folder.number_of_sub_messages):
            try:
                msg = folder.get_sub_message(i)
                subject = msg.subject or ""
                sender = msg.sender_name or ""
                sent_time = str(msg.delivery_time) if msg.delivery_time else ""

                body = ""
                try:
                    plain = msg.plain_text_body
                    if plain:
                        body = plain.decode("utf-8", errors="replace") if isinstance(plain, bytes) else plain
                except Exception:
                    pass

                if not body.strip():
                    try:
                        html = msg.html_body
                        if html:
                            html_str = html.decode("utf-8", errors="replace") if isinstance(html, bytes) else html
                            body = _strip_html(html_str)
                    except Exception:
                        pass

                if body.strip():
                    header = (
                        f"--- [{folder_name}] "
                        f"Subject: {subject} | "
                        f"From: {sender} | "
                        f"Date: {sent_time} ---"
                    )
                    results.append(header)
                    results.append(body.strip())
                    results.append("")

            except Exception as msg_exc:
                logger.debug("Failed to extract PST message %d: %s", i, msg_exc)

        for j in range(folder.number_of_sub_folders):
            try:
                results.extend(_walk_folder(folder.get_sub_folder(j)))
            except Exception as folder_exc:
                logger.debug("Failed to recurse PST sub-folder: %s", folder_exc)

        return results

    try:
        # pypff requires a file path — write to a temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pst") as tmp:
            tmp.write(data)
            tmp_path = tmp.name

        pst = pypff.file()
        parts: list[str] = []
        try:
            pst.open(tmp_path)
            root = pst.get_root_folder()
            for i in range(root.number_of_sub_folders):
                try:
                    parts.extend(_walk_folder(root.get_sub_folder(i)))
                except Exception as fe:
                    logger.debug("PST top-level folder error: %s", fe)
        finally:
            pst.close()
            os.unlink(tmp_path)

        if not parts:
            return "[PST file contained no readable message text]"

        return "\n".join(parts)

    except Exception as exc:
        raise ValueError(f".pst extraction failed: {exc}") from exc


# ---------------------------------------------------------------------------
# HTML tag stripper (no external deps)
# ---------------------------------------------------------------------------


def _strip_html(html: str) -> str:
    """Remove HTML tags and decode common entities to produce plain text."""
    import re

    # Remove script and style blocks entirely
    html = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html, flags=re.DOTALL | re.IGNORECASE)
    # Replace block-level tags with newlines for readability
    html = re.sub(r"<(br|p|div|tr|li|h[1-6])[^>]*>", "\n", html, flags=re.IGNORECASE)
    # Remove all remaining tags
    html = re.sub(r"<[^>]+>", "", html)
    # Decode common HTML entities
    entities = {
        "&amp;": "&", "&lt;": "<", "&gt;": ">",
        "&quot;": '"', "&#39;": "'", "&nbsp;": " ",
        "&apos;": "'",
    }
    for entity, char in entities.items():
        html = html.replace(entity, char)
    # Collapse excessive whitespace
    html = re.sub(r"\n{3,}", "\n\n", html)
    return html.strip()
